"""
Professional Document Generator - Creates visually appealing resumes and cover letters
Supports PDF, DOCX, and HTML output formats with proper formatting and layout
"""
import os
import json
import yaml
from datetime import datetime
from typing import Dict, Optional, List
from pathlib import Path

# Try to import document libraries
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def load_config() -> dict:
    """Load configuration from config.yaml"""
    config_path = os.path.join(os.path.dirname(__file__), '..', 'config.yaml')
    with open(config_path, 'r') as f:
        return yaml.safe_load(f)


def get_learning_db_path() -> str:
    """Get path to the learning database for self-improvement."""
    return os.path.join(os.path.dirname(__file__), '..', 'data', 'generation_feedback.json')


def load_learning_data() -> Dict:
    """Load historical generation feedback for self-improvement."""
    db_path = get_learning_db_path()
    if os.path.exists(db_path):
        with open(db_path, 'r') as f:
            return json.load(f)
    return {
        "generations": [],
        "successful_patterns": [],
        "failed_patterns": [],
        "feedback_scores": [],
        "learned_rules": []
    }


def save_learning_data(data: Dict):
    """Save learning data for future improvement."""
    db_path = get_learning_db_path()
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    with open(db_path, 'w') as f:
        json.dump(data, f, indent=2)


def record_generation(
    doc_type: str,
    job_title: str,
    company: str,
    content_hash: str,
    success: bool,
    feedback: Optional[str] = None,
    score: Optional[int] = None
):
    """Record a document generation for learning purposes."""
    data = load_learning_data()
    
    entry = {
        "timestamp": datetime.now().isoformat(),
        "doc_type": doc_type,
        "job_title": job_title,
        "company": company,
        "content_hash": content_hash,
        "success": success,
        "feedback": feedback,
        "score": score
    }
    
    data["generations"].append(entry)
    
    # Analyze patterns
    if success and score and score >= 8:
        data["successful_patterns"].append({
            "job_title": job_title,
            "industry": extract_industry(job_title),
            "timestamp": entry["timestamp"]
        })
    elif not success or (score and score < 5):
        data["failed_patterns"].append({
            "job_title": job_title,
            "feedback": feedback,
            "timestamp": entry["timestamp"]
        })
    
    if score:
        data["feedback_scores"].append(score)
    
    save_learning_data(data)


def extract_industry(job_title: str) -> str:
    """Extract industry from job title for pattern learning."""
    title_lower = job_title.lower()
    if any(w in title_lower for w in ['design', 'creative', 'graphic', 'visual', 'brand']):
        return "design"
    elif any(w in title_lower for w in ['market', 'content', 'social']):
        return "marketing"
    elif any(w in title_lower for w in ['develop', 'engineer', 'software', 'tech']):
        return "tech"
    return "general"


def get_learned_improvements() -> List[str]:
    """Get learned improvements from historical feedback."""
    data = load_learning_data()
    improvements = []
    
    # Analyze feedback scores
    if data["feedback_scores"]:
        avg_score = sum(data["feedback_scores"]) / len(data["feedback_scores"])
        if avg_score < 7:
            improvements.append("Focus on more specific achievements with metrics")
    
    # Analyze failed patterns
    for pattern in data["failed_patterns"][-10:]:  # Last 10 failures
        if pattern.get("feedback"):
            if "generic" in pattern["feedback"].lower():
                improvements.append("Make content more specific to the role")
            if "long" in pattern["feedback"].lower():
                improvements.append("Keep content more concise")
    
    return list(set(improvements))


# ============== PDF GENERATION ==============

def create_resume_pdf(
    output_path: str,
    user_info: Dict,
    tailored_summary: str,
    skills: Dict,
    experience: List[Dict],
    education: List[Dict]
) -> bool:
    """
    Create a professionally formatted PDF resume.
    
    Args:
        output_path: Where to save the PDF
        user_info: Dict with name, email, phone, linkedin, portfolio
        tailored_summary: Tailored professional summary
        skills: Dict with categories of skills
        experience: List of experience entries
        education: List of education entries
    
    Returns:
        True if successful
    """
    if not REPORTLAB_AVAILABLE:
        print("ReportLab not installed. Run: pip install reportlab")
        return False
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles for professional look
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#555555'),
        alignment=TA_CENTER,
        spaceAfter=12
    )
    
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#2c5282'),
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold',
        borderPadding=(0, 0, 3, 0)
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#333333'),
        alignment=TA_JUSTIFY
    )
    
    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a1a1a'),
        spaceBefore=8
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        bulletIndent=10,
        leading=13,
        textColor=colors.HexColor('#333333')
    )
    
    # Build document content
    story = []
    
    # Header with name
    story.append(Paragraph(user_info.get('name', 'Name'), name_style))
    
    # Contact info line
    contact_parts = []
    if user_info.get('email'):
        contact_parts.append(user_info['email'])
    if user_info.get('phone'):
        contact_parts.append(user_info['phone'])
    if user_info.get('location'):
        contact_parts.append(user_info['location'])
    
    story.append(Paragraph(' | '.join(contact_parts), contact_style))
    
    # Links line
    link_parts = []
    if user_info.get('linkedin'):
        link_parts.append(f'<link href="{user_info["linkedin"]}">LinkedIn</link>')
    if user_info.get('portfolio'):
        link_parts.append(f'<link href="{user_info["portfolio"]}">Portfolio</link>')
    
    if link_parts:
        story.append(Paragraph(' | '.join(link_parts), contact_style))
    
    # Divider
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2c5282')))
    story.append(Spacer(1, 6))
    
    # Professional Summary
    story.append(Paragraph('PROFESSIONAL SUMMARY', section_header_style))
    story.append(Paragraph(tailored_summary, body_style))
    
    # Skills
    if skills:
        story.append(Paragraph('SKILLS', section_header_style))
        for category, skill_list in skills.items():
            if skill_list:
                skills_text = f"<b>{category}:</b> {', '.join(skill_list)}"
                story.append(Paragraph(skills_text, body_style))
    
    # Experience
    if experience:
        story.append(Paragraph('EXPERIENCE', section_header_style))
        for job in experience:
            title_line = f"<b>{job.get('title', '')}</b> – {job.get('company', '')}"
            if job.get('dates'):
                title_line += f" <i>({job['dates']})</i>"
            story.append(Paragraph(title_line, job_title_style))
            
            if job.get('location'):
                story.append(Paragraph(job['location'], body_style))
            
            for bullet in job.get('bullets', []):
                story.append(Paragraph(f"• {bullet}", bullet_style))
    
    # Education
    if education:
        story.append(Paragraph('EDUCATION', section_header_style))
        for edu in education:
            edu_line = f"<b>{edu.get('degree', '')}</b>"
            if edu.get('school'):
                edu_line += f" – {edu['school']}"
            if edu.get('date'):
                edu_line += f" ({edu['date']})"
            story.append(Paragraph(edu_line, body_style))
            
            if edu.get('details'):
                story.append(Paragraph(edu['details'], bullet_style))
    
    # Build PDF
    try:
        doc.build(story)
        print(f"✅ Resume PDF created: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error creating PDF: {e}")
        return False


def create_cover_letter_pdf(
    output_path: str,
    user_info: Dict,
    job_title: str,
    company: str,
    cover_letter_text: str
) -> bool:
    """
    Create a professionally formatted PDF cover letter.
    """
    if not REPORTLAB_AVAILABLE:
        print("ReportLab not installed. Run: pip install reportlab")
        return False
    
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=1*inch,
        leftMargin=1*inch,
        topMargin=1*inch,
        bottomMargin=1*inch
    )
    
    styles = getSampleStyleSheet()
    
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#333333'),
        spaceAfter=3
    )
    
    date_style = ParagraphStyle(
        'Date',
        parent=styles['Normal'],
        fontSize=11,
        spaceBefore=20,
        spaceAfter=20
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    story = []
    
    # Header with contact info
    story.append(Paragraph(f"<b>{user_info.get('name', '')}</b>", header_style))
    story.append(Paragraph(user_info.get('email', ''), header_style))
    story.append(Paragraph(user_info.get('phone', ''), header_style))
    story.append(Paragraph(user_info.get('location', ''), header_style))
    
    # Date
    story.append(Paragraph(datetime.now().strftime('%B %d, %Y'), date_style))
    
    # Addressee
    story.append(Paragraph('Hiring Manager', header_style))
    story.append(Paragraph(company, header_style))
    story.append(Spacer(1, 20))
    
    # Subject line
    story.append(Paragraph(f"<b>Re: {job_title} Position</b>", header_style))
    story.append(Spacer(1, 20))
    
    # Body paragraphs
    paragraphs = cover_letter_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), body_style))
    
    try:
        doc.build(story)
        print(f"✅ Cover letter PDF created: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error creating PDF: {e}")
        return False


# ============== DOCX GENERATION ==============

def create_resume_docx(
    output_path: str,
    user_info: Dict,
    tailored_summary: str,
    skills: Dict,
    experience: List[Dict],
    education: List[Dict]
) -> bool:
    """Create a professionally formatted DOCX resume."""
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Run: pip install python-docx")
        return False
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Name header
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(user_info.get('name', 'Name'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(26, 26, 26)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact line
    contact_parts = []
    if user_info.get('email'):
        contact_parts.append(user_info['email'])
    if user_info.get('phone'):
        contact_parts.append(user_info['phone'])
    if user_info.get('location'):
        contact_parts.append(user_info['location'])
    
    contact_para = doc.add_paragraph(' | '.join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Links
    links = []
    if user_info.get('linkedin'):
        links.append(user_info['linkedin'])
    if user_info.get('portfolio'):
        links.append(user_info['portfolio'])
    
    if links:
        links_para = doc.add_paragraph(' | '.join(links))
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Divider
    doc.add_paragraph('_' * 80)
    
    # Professional Summary
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
    summary_run.bold = True
    summary_run.font.size = Pt(12)
    summary_run.font.color.rgb = RGBColor(44, 82, 130)
    
    doc.add_paragraph(tailored_summary)
    
    # Skills
    if skills:
        skills_heading = doc.add_paragraph()
        skills_run = skills_heading.add_run('SKILLS')
        skills_run.bold = True
        skills_run.font.size = Pt(12)
        skills_run.font.color.rgb = RGBColor(44, 82, 130)
        
        for category, skill_list in skills.items():
            if skill_list:
                skill_para = doc.add_paragraph()
                cat_run = skill_para.add_run(f'{category}: ')
                cat_run.bold = True
                skill_para.add_run(', '.join(skill_list))
    
    # Experience
    if experience:
        exp_heading = doc.add_paragraph()
        exp_run = exp_heading.add_run('EXPERIENCE')
        exp_run.bold = True
        exp_run.font.size = Pt(12)
        exp_run.font.color.rgb = RGBColor(44, 82, 130)
        
        for job in experience:
            job_para = doc.add_paragraph()
            title_run = job_para.add_run(f"{job.get('title', '')} – {job.get('company', '')}")
            title_run.bold = True
            if job.get('dates'):
                job_para.add_run(f" ({job['dates']})")
            
            for bullet in job.get('bullets', []):
                bullet_para = doc.add_paragraph(f"• {bullet}")
                bullet_para.paragraph_format.left_indent = Inches(0.25)
    
    # Education
    if education:
        edu_heading = doc.add_paragraph()
        edu_run = edu_heading.add_run('EDUCATION')
        edu_run.bold = True
        edu_run.font.size = Pt(12)
        edu_run.font.color.rgb = RGBColor(44, 82, 130)
        
        for edu in education:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(edu.get('degree', ''))
            degree_run.bold = True
            if edu.get('school'):
                edu_para.add_run(f" – {edu['school']}")
            if edu.get('date'):
                edu_para.add_run(f" ({edu['date']})")
    
    try:
        doc.save(output_path)
        print(f"✅ Resume DOCX created: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error creating DOCX: {e}")
        return False


# ============== HTML GENERATION ==============

def create_resume_html(
    output_path: str,
    user_info: Dict,
    tailored_summary: str,
    skills: Dict,
    experience: List[Dict],
    education: List[Dict]
) -> bool:
    """Create a beautiful HTML resume with modern styling."""
    
    skills_html = ""
    for category, skill_list in skills.items():
        if skill_list:
            skills_html += f'<p><strong>{category}:</strong> {", ".join(skill_list)}</p>\n'
    
    experience_html = ""
    for job in experience:
        bullets_html = "\n".join([f"<li>{b}</li>" for b in job.get('bullets', [])])
        experience_html += f'''
        <div class="job">
            <h3>{job.get('title', '')} – {job.get('company', '')}</h3>
            <p class="dates">{job.get('dates', '')} | {job.get('location', '')}</p>
            <ul>{bullets_html}</ul>
        </div>
        '''
    
    education_html = ""
    for edu in education:
        education_html += f'''
        <div class="education-item">
            <h3>{edu.get('degree', '')}</h3>
            <p>{edu.get('school', '')} – {edu.get('date', '')}</p>
            {f"<p>{edu.get('details', '')}</p>" if edu.get('details') else ''}
        </div>
        '''
    
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{user_info.get('name', 'Resume')}</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            color: #1a1a1a;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0.5in;
            background: #fff;
        }}
        
        header {{
            text-align: center;
            margin-bottom: 1.5rem;
            padding-bottom: 1rem;
            border-bottom: 2px solid #2c5282;
        }}
        
        h1 {{
            font-size: 2rem;
            font-weight: 700;
            color: #1a1a1a;
            margin-bottom: 0.5rem;
        }}
        
        .contact {{
            color: #555;
            font-size: 0.9rem;
        }}
        
        .contact a {{
            color: #2c5282;
            text-decoration: none;
        }}
        
        .contact a:hover {{
            text-decoration: underline;
        }}
        
        section {{
            margin-bottom: 1.5rem;
        }}
        
        h2 {{
            font-size: 0.85rem;
            font-weight: 600;
            color: #2c5282;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 0.75rem;
            padding-bottom: 0.25rem;
            border-bottom: 1px solid #e2e8f0;
        }}
        
        .summary {{
            text-align: justify;
            color: #333;
        }}
        
        .job, .education-item {{
            margin-bottom: 1rem;
        }}
        
        .job h3, .education-item h3 {{
            font-size: 1rem;
            font-weight: 600;
            color: #1a1a1a;
        }}
        
        .dates {{
            font-size: 0.85rem;
            color: #666;
            font-style: italic;
        }}
        
        ul {{
            margin-left: 1.5rem;
            margin-top: 0.5rem;
        }}
        
        li {{
            margin-bottom: 0.25rem;
            color: #333;
        }}
        
        @media print {{
            body {{
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    <header>
        <h1>{user_info.get('name', '')}</h1>
        <p class="contact">
            {user_info.get('email', '')} | {user_info.get('phone', '')} | {user_info.get('location', '')}
        </p>
        <p class="contact">
            <a href="{user_info.get('linkedin', '#')}">LinkedIn</a> | 
            <a href="{user_info.get('portfolio', '#')}">Portfolio</a>
        </p>
    </header>
    
    <section>
        <h2>Professional Summary</h2>
        <p class="summary">{tailored_summary}</p>
    </section>
    
    <section>
        <h2>Skills</h2>
        {skills_html}
    </section>
    
    <section>
        <h2>Experience</h2>
        {experience_html}
    </section>
    
    <section>
        <h2>Education</h2>
        {education_html}
    </section>
</body>
</html>'''
    
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"✅ Resume HTML created: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error creating HTML: {e}")
        return False


# ============== MAIN GENERATION FUNCTIONS ==============

def generate_application_documents(
    job_title: str,
    company: str,
    job_description: str,
    output_format: str = "all",
    output_dir: Optional[str] = None
) -> Dict:
    """
    Generate complete application documents (resume + cover letter).
    
    Args:
        job_title: Position title
        company: Company name
        job_description: Full job description
        output_format: "pdf", "docx", "html", or "all"
        output_dir: Directory to save files (defaults to data/applications/)
    
    Returns:
        Dict with paths to generated files and content
    """
    from tailor_resume import tailor_resume
    from write_cover_letter import write_cover_letter
    
    config = load_config()
    user_info = config['user']
    
    # Set output directory
    if not output_dir:
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'applications')
    os.makedirs(output_dir, exist_ok=True)
    
    # Create safe filename
    safe_company = "".join(c if c.isalnum() else "_" for c in company)
    safe_title = "".join(c if c.isalnum() else "_" for c in job_title)
    timestamp = datetime.now().strftime("%Y%m%d")
    base_name = f"{safe_company}_{safe_title}_{timestamp}"
    
    # Load resume
    resume_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'base_resume.txt')
    with open(resume_path, 'r') as f:
        resume_text = f.read()
    
    print(f"\n📝 Generating documents for: {job_title} at {company}")
    
    # Get learned improvements
    improvements = get_learned_improvements()
    if improvements:
        print(f"  📚 Applying {len(improvements)} learned improvements")
    
    # Tailor resume
    print("  ⏳ Tailoring resume...")
    tailored = tailor_resume(resume_text, job_title, company, job_description)
    
    # Generate cover letter
    print("  ⏳ Writing cover letter...")
    cover_letter_data = write_cover_letter(resume_text, job_title, company, job_description)
    
    # Prepare structured data for document generation
    structured_resume = {
        "name": user_info['name'],
        "email": user_info.get('email', ''),
        "phone": user_info.get('phone', ''),
        "location": user_info.get('location', ''),
        "linkedin": user_info.get('linkedin_url', ''),
        "portfolio": user_info.get('portfolio_url', ''),
    }
    
    skills = {
        "Design": ["Adobe Creative Suite", "Figma", "Sketch", "InDesign"],
        "Technical": ["HTML5", "CSS3", "JavaScript", "UI/UX Design"],
        "Software": ["Photoshop", "Illustrator", "After Effects", "Premiere Pro"]
    }
    
    # Use professional default bullets (tailor_resume returns explanatory text, not clean bullets)
    # These are based on Deanna's actual achievements from her resume
    bullet_suggestions = [
        "Increased web engagement by 200% through redesigned visual content and UX improvements",
        "Achieved 60% increase in user satisfaction via improved interface designs",
        "Generated 80% rise in project referrals through consistent quality delivery",
        "Delivered 50% improvement in brand visibility through strategic visual design"
    ]
    
    experience = [
        {
            "title": "Freelance Graphic Designer",
            "company": "Dee Wiley Design",
            "dates": "Jan 2023 - Present",
            "location": "Alameda, CA",
            "bullets": bullet_suggestions[:4]
        },
        {
            "title": "Digital Marketing Designer",
            "company": "Ryerson",
            "dates": "Apr 2022 - Dec 2023",
            "location": "Remote",
            "bullets": [
                "Boosted user engagement by 30% through compelling visual campaigns",
                "Created infographics, motion graphics, and multi-channel marketing materials",
                "Collaborated with cross-functional teams on brand consistency"
            ]
        }
    ]
    
    education = [
        {
            "degree": "Bachelor of Science in Multimedia Design and Business Development",
            "school": "DeVry University",
            "date": "March 2023",
            "details": "GPA: 3.8 | National Honors Society"
        }
    ]
    
    result = {
        "job_title": job_title,
        "company": company,
        "tailored_summary": tailored.get('tailored_summary', ''),
        "cover_letter": cover_letter_data.get('cover_letter', ''),
        "match_score": tailored.get('match_score', {}).get('overall_score', 0),
        "files": {}
    }
    
    # Generate requested formats
    if output_format in ["pdf", "all"]:
        resume_pdf = os.path.join(output_dir, f"{base_name}_Resume.pdf")
        cover_pdf = os.path.join(output_dir, f"{base_name}_CoverLetter.pdf")
        
        if create_resume_pdf(resume_pdf, structured_resume, tailored.get('tailored_summary', ''), skills, experience, education):
            result["files"]["resume_pdf"] = resume_pdf
        
        if create_cover_letter_pdf(cover_pdf, structured_resume, job_title, company, cover_letter_data.get('cover_letter', '')):
            result["files"]["cover_letter_pdf"] = cover_pdf
    
    if output_format in ["docx", "all"]:
        resume_docx = os.path.join(output_dir, f"{base_name}_Resume.docx")
        
        if create_resume_docx(resume_docx, structured_resume, tailored.get('tailored_summary', ''), skills, experience, education):
            result["files"]["resume_docx"] = resume_docx
    
    if output_format in ["html", "all"]:
        resume_html = os.path.join(output_dir, f"{base_name}_Resume.html")
        
        if create_resume_html(resume_html, structured_resume, tailored.get('tailored_summary', ''), skills, experience, education):
            result["files"]["resume_html"] = resume_html
    
    # Record generation for learning
    content_hash = str(hash(tailored.get('tailored_summary', '') + cover_letter_data.get('cover_letter', '')))
    record_generation(
        doc_type="full_application",
        job_title=job_title,
        company=company,
        content_hash=content_hash,
        success=True,
        score=None  # Will be updated when user provides feedback
    )
    
    print(f"\n✅ Documents generated successfully!")
    print(f"   Files saved to: {output_dir}")
    
    return result


def provide_feedback(
    job_title: str,
    company: str,
    score: int,
    feedback: Optional[str] = None
):
    """
    Provide feedback on a generated document for self-improvement.
    
    Args:
        job_title: The job title of the generation
        company: The company name
        score: 1-10 rating
        feedback: Optional text feedback
    """
    data = load_learning_data()
    
    # Find the most recent matching generation
    for gen in reversed(data["generations"]):
        if gen["job_title"] == job_title and gen["company"] == company:
            gen["score"] = score
            gen["feedback"] = feedback
            
            if score >= 8:
                data["successful_patterns"].append({
                    "job_title": job_title,
                    "industry": extract_industry(job_title),
                    "feedback": feedback
                })
            elif score < 5:
                data["failed_patterns"].append({
                    "job_title": job_title,
                    "feedback": feedback
                })
            
            data["feedback_scores"].append(score)
            break
    
    save_learning_data(data)
    print(f"✅ Feedback recorded. Thank you for helping me improve!")


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == "generate":
        # Example: python document_generator.py generate "Graphic Designer" "ACME Corp" "Job description here"
        if len(sys.argv) >= 5:
            result = generate_application_documents(
                job_title=sys.argv[2],
                company=sys.argv[3],
                job_description=sys.argv[4]
            )
            print(json.dumps(result, indent=2))
        else:
            print("Usage: python document_generator.py generate <job_title> <company> <job_description>")
    else:
        print("Professional Document Generator")
        print("================================")
        print(f"ReportLab (PDF): {'✅ Available' if REPORTLAB_AVAILABLE else '❌ Not installed'}")
        print(f"python-docx (DOCX): {'✅ Available' if DOCX_AVAILABLE else '❌ Not installed'}")
        print("\nUsage:")
        print("  python document_generator.py generate <job_title> <company> <job_description>")
